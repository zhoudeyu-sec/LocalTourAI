"""
知识库批量导入脚本
从 docx 文件中读取数据，导入到 MySQL 数据库
"""

import os
import re
from database import SessionLocal
from models import KnowledgeBase


DATA_DIR = "data"


def import_from_docx():
    """从 docx 文件读取景点数据"""
    try:
        from docx import Document
    except ImportError:
        print("❌ 缺少 python-docx 库，请运行: pip install python-docx")
        return []

    docx_path = os.path.join(DATA_DIR, "灵山胜境 景点结构化数据集.docx")
    if not os.path.exists(docx_path):
        print(f"⚠️ 文件不存在: {docx_path}")
        return []

    doc = Document(docx_path)
    items = []
    
    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows[1:]:  # 跳过表头
            cells = row.cells
            if len(cells) >= 4:
                # 表格列：景区名称、景点ID、景点名称、具体位置...
                spot_id = cells[1].text.strip() if len(cells) > 1 else ""
                name = cells[2].text.strip() if len(cells) > 2 else ""
                location = cells[3].text.strip() if len(cells) > 3 else ""
                desc = cells[7].text.strip() if len(cells) > 7 else ""
                
                if name and len(name) > 2:
                    content = f"位置：{location}\n详细介绍：{desc}" if desc else location
                    items.append({
                        "title": f"{spot_id} {name}" if spot_id else name,
                        "content": content[:3000],
                        "category": "景点介绍"
                    })
    
    print(f"📄 从表格中提取到 {len(items)} 条数据")
    return items


def import_from_guide():
    """从导游指南读取数据"""
    docx_path = os.path.join(DATA_DIR, "灵山胜境：历史、文化、景点特色与个性化游览指南.docx")
    if not os.path.exists(docx_path):
        print(f"⚠️ 文件不存在: {docx_path}")
        return []

    try:
        from docx import Document
    except ImportError:
        return []

    doc = Document(docx_path)
    items = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 5:
            continue
        
        # 识别标题（以 ## 开头或加粗样式）
        if text.startswith('##') or (len(text) < 30 and ('景点' in text or '文化' in text or '历史' in text)):
            title = text.replace('##', '').strip()
            # 获取后续段落作为内容
            content_parts = []
            for next_para in doc.paragraphs[doc.paragraphs.index(para)+1:]:
                next_text = next_para.text.strip()
                if next_text.startswith('##') or (len(next_text) < 30 and '景点' in next_text):
                    break
                if next_text:
                    content_parts.append(next_text)
            
            if content_parts:
                items.append({
                    "title": title[:100],
                    "content": " ".join(content_parts)[:3000],
                    "category": "历史文化"
                })
    
    print(f"📄 从指南中提取到 {len(items)} 条数据")
    return items


def import_to_db(items):
    """导入到数据库"""
    if not items:
        print("⚠️ 没有数据可导入")
        return 0

    print(f"\n📚 准备导入 {len(items)} 条数据...")
    
    db = SessionLocal()
    count = 0
    
    try:
        for item in items:
            title = item.get("title", "")
            content = item.get("content", "")
            category = item.get("category", "景点介绍")

            if not title or not content:
                continue

            # 检查是否已存在
            existing = db.query(KnowledgeBase).filter(KnowledgeBase.title == title).first()
            if existing:
                print(f"⏭️ 已存在: {title[:30]}...")
                continue

            try:
                kb = KnowledgeBase(
                    title=title[:255],
                    content=content[:5000],
                    category=category
                )
                db.add(kb)
                db.commit()
                count += 1
                print(f"✅ 导入: {title[:40]}...")
            except Exception as e:
                db.rollback()
                print(f"❌ 失败 {title[:30]}: {e}")
    finally:
        db.close()

    return count


def main():
    print("=" * 50)
    print("灵山景区知识库批量导入工具")
    print("=" * 50)

    # 导入数据
    items = import_from_docx()
    if not items:
        items = import_from_guide()

    if items:
        count = import_to_db(items)
        print(f"\n🎉 导入完成！共导入 {count} 条知识")
    else:
        print("⚠️ 未找到可导入的数据")

    # 显示导入结果
    db = SessionLocal()
    total = db.query(KnowledgeBase).count()
    db.close()
    print(f"\n📊 知识库当前共有 {total} 条数据")


if __name__ == "__main__":
    main()